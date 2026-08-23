"""
Assignment workflow with VERSIONING: every upload/generate/edit/regenerate
creates a NEW AssignmentVersionORM row - nothing is ever overwritten.
Approving picks ONE version to be "the" sent version; approving an OLDER
version than the latest is how HR rolls back an accidental approval.

Only /generate and /regenerate use the LLM (1 call each). Everything else
here is plain CRUD + a fixed-template email send.
"""

from __future__ import annotations

import io
import tempfile
import uuid
from datetime import datetime, timedelta, timezone

from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.schemas import (
    ApproveAssignmentRequest,
    AssignmentResponse,
    AssignmentSendResult,
    AssignmentVersionResponse,
    CreateAssignmentLinkRequest,
    EditAssignmentRequest,
    ExtendDeadlineRequest,
    GenerateCampaignAssignmentRequest,
)
from app.core.config import get_settings
from app.core.container import get_email_sender_dep, get_generate_assignment_use_case
from app.core.exceptions import DocumentParsingError, LLMEvaluationError
from app.core.logging import get_logger
from app.infrastructure.db.database import get_session
from app.infrastructure.db.models import AssignmentORM, AssignmentVersionORM, CampaignEventORM, CampaignORM, CandidateORM
from app.infrastructure.email.templates import ASSIGNMENT_EMAIL_BODY, ASSIGNMENT_EMAIL_SUBJECT
from app.infrastructure.parsers.document_factory import get_parser_for
from app.utils.pdf import render_paginated_text_pdf

router = APIRouter(prefix="/api/v1", tags=["assignments"])
logger = get_logger(__name__)

_ALLOWED_SUFFIXES = {".docx", ".pdf"}


def _as_utc(dt: datetime | None) -> datetime | None:
    if dt is None:
        return None
    return dt if dt.tzinfo is not None else dt.replace(tzinfo=timezone.utc)


async def _log_event(session: AsyncSession, campaign_id: str, event_type: str, message: str) -> None:
    session.add(CampaignEventORM(campaign_id=campaign_id, event_type=event_type, message=message))


async def _next_version_number(session: AsyncSession, assignment_id: str) -> int:
    result = await session.execute(select(AssignmentVersionORM).where(AssignmentVersionORM.assignment_id == assignment_id))
    existing = result.scalars().all()
    return (max((v.version_number for v in existing), default=0)) + 1


async def _to_response(session: AsyncSession, a: AssignmentORM) -> AssignmentResponse:
    result = await session.execute(select(AssignmentVersionORM).where(AssignmentVersionORM.assignment_id == a.id).order_by(AssignmentVersionORM.version_number.desc()))
    versions = result.scalars().all()
    current = versions[0]  # latest by version_number - what HR is currently looking at

    return AssignmentResponse(
        id=a.id, campaign_id=a.campaign_id, title=a.title, status=a.status,
        deadline=_as_utc(a.deadline), sent_at=_as_utc(a.sent_at), created_at=_as_utc(a.created_at),
        current_version=AssignmentVersionResponse(
            id=current.id, version_number=current.version_number, question_paper_text=current.question_paper_text,
            source=current.source, label=current.label, google_drive_link=current.google_drive_link, created_at=_as_utc(current.created_at),
            is_approved=(current.id == a.approved_version_id),
        ),
        version_count=len(versions),
    )


@router.post("/campaigns/{campaign_id}/assignments/upload", response_model=AssignmentResponse)
async def upload_assignment(
    campaign_id: str,
    title: str = Form(...),
    file: UploadFile = File(..., description="Assignment file (.docx/.pdf)"),
    session: AsyncSession = Depends(get_session),
) -> AssignmentResponse:
    campaign = await session.get(CampaignORM, campaign_id)
    if campaign is None:
        raise HTTPException(status_code=404, detail=f"Campaign '{campaign_id}' not found.")

    suffix = "." + file.filename.rsplit(".", 1)[-1].lower() if file.filename and "." in file.filename else ""
    if suffix not in _ALLOWED_SUFFIXES:
        raise HTTPException(status_code=415, detail=f"Unsupported file type '{suffix}'. Allowed: .docx, .pdf")

    with tempfile.TemporaryDirectory() as tmp_dir:
        path = f"{tmp_dir}/{file.filename}"
        with open(path, "wb") as f:
            f.write(await file.read())
        try:
            question_paper_text = await get_parser_for(path).extract_text(path)
        except DocumentParsingError as e:
            raise HTTPException(status_code=422, detail=str(e)) from e

    assignment = AssignmentORM(campaign_id=campaign_id, title=title, status="draft")
    session.add(assignment)
    await session.flush()
    version = AssignmentVersionORM(assignment_id=assignment.id, version_number=1, question_paper_text=question_paper_text, source="uploaded", label="Uploaded")
    session.add(version)
    await session.commit()
    await session.refresh(assignment)
    return await _to_response(session, assignment)


@router.post("/campaigns/{campaign_id}/assignments/link", response_model=AssignmentResponse)
async def link_assignment(campaign_id: str, body: CreateAssignmentLinkRequest, session: AsyncSession = Depends(get_session)) -> AssignmentResponse:
    """The other way (besides upload/generate) HR can create an
    assignment: point at an existing Google Drive file instead of
    uploading one. Same versioning rules apply."""

    campaign = await session.get(CampaignORM, campaign_id)
    if campaign is None:
        raise HTTPException(status_code=404, detail=f"Campaign '{campaign_id}' not found.")

    assignment = AssignmentORM(campaign_id=campaign_id, title=body.title, status="draft")
    session.add(assignment)
    await session.flush()
    version = AssignmentVersionORM(
        assignment_id=assignment.id, version_number=1, question_paper_text="", source="link",
        label="Google Drive Link", google_drive_link=body.google_drive_link,
    )
    session.add(version)
    await session.commit()
    await session.refresh(assignment)
    return await _to_response(session, assignment)


@router.post("/campaigns/{campaign_id}/assignments/generate", response_model=AssignmentResponse)
async def generate_campaign_assignment(
    campaign_id: str, body: GenerateCampaignAssignmentRequest, session: AsyncSession = Depends(get_session)
) -> AssignmentResponse:
    """1 LLM call. Creates the assignment AND its version 1."""

    campaign = await session.get(CampaignORM, campaign_id)
    if campaign is None:
        raise HTTPException(status_code=404, detail=f"Campaign '{campaign_id}' not found.")

    use_case = get_generate_assignment_use_case()
    try:
        question_paper_text = await use_case.execute(
            brief=body.brief, job_description_text=campaign.job_description_text,
            num_questions=body.num_questions, num_sections=body.num_sections, difficulty=body.difficulty,
        )
    except LLMEvaluationError as e:
        raise HTTPException(status_code=502, detail=f"Assignment generation failed: {e}") from e

    assignment = AssignmentORM(campaign_id=campaign_id, title=body.title, status="draft")
    session.add(assignment)
    await session.flush()
    version = AssignmentVersionORM(assignment_id=assignment.id, version_number=1, question_paper_text=question_paper_text, source="generated", label="Generated")
    session.add(version)
    await _log_event(session, campaign_id, "assignment_generated", f"Assignment '{body.title}' generated (v1).")
    await session.commit()
    await session.refresh(assignment)
    return await _to_response(session, assignment)


@router.get("/assignments/{assignment_id}", response_model=AssignmentResponse)
async def get_assignment(assignment_id: str, session: AsyncSession = Depends(get_session)) -> AssignmentResponse:
    assignment = await session.get(AssignmentORM, assignment_id)
    if assignment is None:
        raise HTTPException(status_code=404, detail=f"Assignment '{assignment_id}' not found.")
    return await _to_response(session, assignment)


@router.get("/assignments/{assignment_id}/versions", response_model=list[AssignmentVersionResponse])
async def list_assignment_versions(assignment_id: str, session: AsyncSession = Depends(get_session)) -> list[AssignmentVersionResponse]:
    """The version history view - 'Version 1 Generated 2:30 PM, Version 2
    Generated 2:35 PM, ...'."""

    assignment = await session.get(AssignmentORM, assignment_id)
    if assignment is None:
        raise HTTPException(status_code=404, detail=f"Assignment '{assignment_id}' not found.")

    result = await session.execute(select(AssignmentVersionORM).where(AssignmentVersionORM.assignment_id == assignment_id).order_by(AssignmentVersionORM.version_number))
    versions = result.scalars().all()
    return [
        AssignmentVersionResponse(id=v.id, version_number=v.version_number, question_paper_text=v.question_paper_text, source=v.source, label=v.label, google_drive_link=v.google_drive_link, created_at=v.created_at, is_approved=(v.id == assignment.approved_version_id))
        for v in versions
    ]


@router.patch("/assignments/{assignment_id}", response_model=AssignmentResponse)
async def edit_assignment(assignment_id: str, body: EditAssignmentRequest, session: AsyncSession = Depends(get_session)) -> AssignmentResponse:
    """HR's editor 'Save' - creates a NEW version (never overwrites)."""

    assignment = await session.get(AssignmentORM, assignment_id)
    if assignment is None:
        raise HTTPException(status_code=404, detail=f"Assignment '{assignment_id}' not found.")
    if assignment.status != "draft":
        raise HTTPException(status_code=409, detail="Only draft assignments can be edited.")

    next_number = await _next_version_number(session, assignment_id)
    session.add(AssignmentVersionORM(assignment_id=assignment_id, version_number=next_number, question_paper_text=body.question_paper_text, source="edited", label="Edited by HR"))
    await session.commit()
    await session.refresh(assignment)
    return await _to_response(session, assignment)


@router.post("/assignments/{assignment_id}/regenerate", response_model=AssignmentResponse)
async def regenerate_assignment(assignment_id: str, body: GenerateCampaignAssignmentRequest, session: AsyncSession = Depends(get_session)) -> AssignmentResponse:
    """1 LLM call.

    LOCKING RULE: an APPROVED assignment can never be edited/replaced in
    place - candidates may already have it. Calling regenerate on an
    approved assignment therefore does NOT touch it at all; instead it
    creates a brand-new, separate draft Assignment (its own version
    history, starting at v1) for the same campaign. This is "start a new
    assignment round", not "change what was already sent".

    Regenerating a DRAFT (not yet approved) instead just adds a new
    version to that same draft, same as before - nothing has been sent
    to candidates yet, so there's nothing to protect."""

    assignment = await session.get(AssignmentORM, assignment_id)
    if assignment is None:
        raise HTTPException(status_code=404, detail=f"Assignment '{assignment_id}' not found.")

    campaign = await session.get(CampaignORM, assignment.campaign_id) if assignment.campaign_id else None
    job_description_text = campaign.job_description_text if campaign else ""

    use_case = get_generate_assignment_use_case()
    try:
        question_paper_text = await use_case.execute(
            brief=body.brief, job_description_text=job_description_text,
            num_questions=body.num_questions, num_sections=body.num_sections, difficulty=body.difficulty,
        )
    except LLMEvaluationError as e:
        raise HTTPException(status_code=502, detail=f"Assignment generation failed: {e}") from e

    if assignment.status == "approved":
        new_assignment = AssignmentORM(campaign_id=assignment.campaign_id, title=body.title, status="draft")
        session.add(new_assignment)
        await session.flush()
        session.add(AssignmentVersionORM(assignment_id=new_assignment.id, version_number=1, question_paper_text=question_paper_text, source="generated", label="Generated (new round)"))
        await session.commit()
        await session.refresh(new_assignment)
        logger.info("Assignment %s is approved (locked) - regenerate created a NEW draft %s instead", assignment_id, new_assignment.id)
        return await _to_response(session, new_assignment)

    next_number = await _next_version_number(session, assignment_id)
    session.add(AssignmentVersionORM(assignment_id=assignment_id, version_number=next_number, question_paper_text=question_paper_text, source="generated", label="Regenerated"))
    assignment.title = body.title
    await session.commit()
    await session.refresh(assignment)
    return await _to_response(session, assignment)


def _render_docx(title: str, question_paper_text: str) -> io.BytesIO:
    """Shared by the HR download route below and the candidate-facing
    download route in routes_submissions.py (imported locally there to
    avoid a module-load-order cycle, same pattern as _to_response)."""

    doc = Document()
    doc.add_heading(title, level=1)
    for line in question_paper_text.split("\n"):
        doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _render_pdf(title: str, question_paper_text: str) -> io.BytesIO:
    """Shared by the HR download route below - renders the question paper
    as a simple paginated PDF via the shared PyMuPDF renderer."""

    return render_paginated_text_pdf(title, question_paper_text.split("\n"))


def candidate_assignment_link(version: AssignmentVersionORM, submission_token: str) -> tuple[str, bool]:
    """Where a candidate's "Open Assignment"/"Download Assignment" should
    point. Returns (link, is_drive_link). Never the HR-only
    /api/v1/assignments/{id}/download route (JWT-gated) - candidates use
    the token-gated /api/v1/submissions/{token}/assignment route instead."""

    if version.google_drive_link:
        return version.google_drive_link, True
    settings = get_settings()
    return f"{settings.backend_public_url}/api/v1/submissions/{submission_token}/assignment", False


@router.get("/assignments/{assignment_id}/download")
async def download_assignment(
    assignment_id: str,
    version_id: str | None = None,
    format: str = Query("docx", pattern="^(docx|pdf)$"),
    session: AsyncSession = Depends(get_session),
) -> StreamingResponse:
    """Downloads a specific version (default: latest) as .docx or .pdf.
    HR-only (this whole router requires a JWT) - candidates use the
    submissions router's token-gated equivalent instead."""

    assignment = await session.get(AssignmentORM, assignment_id)
    if assignment is None:
        raise HTTPException(status_code=404, detail=f"Assignment '{assignment_id}' not found.")

    if version_id:
        version = await session.get(AssignmentVersionORM, version_id)
        if version is None or version.assignment_id != assignment_id:
            raise HTTPException(status_code=404, detail=f"Version '{version_id}' not found for this assignment.")
    else:
        result = await session.execute(select(AssignmentVersionORM).where(AssignmentVersionORM.assignment_id == assignment_id).order_by(AssignmentVersionORM.version_number.desc()).limit(1))
        version = result.scalar_one()

    base_filename = f"{assignment.title.replace(' ', '_')}_v{version.version_number}"
    if format == "pdf":
        buffer = _render_pdf(assignment.title, version.question_paper_text)
        return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{base_filename}.pdf"'})

    buffer = _render_docx(assignment.title, version.question_paper_text)
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{base_filename}.docx"'})


@router.post("/assignments/{assignment_id}/approve", response_model=AssignmentSendResult)
async def approve_assignment(assignment_id: str, body: ApproveAssignmentRequest, session: AsyncSession = Depends(get_session)) -> AssignmentSendResult:
    """Approves a SPECIFIC version (default: latest) - passing an older
    version_id is how HR rolls back an accidental approval. Locks the
    assignment, sets the deadline, and emails every selected candidate
    (Email 2, fixed template, no LLM)."""

    assignment = await session.get(AssignmentORM, assignment_id)
    if assignment is None:
        raise HTTPException(status_code=404, detail=f"Assignment '{assignment_id}' not found.")
    if assignment.campaign_id is None:
        raise HTTPException(status_code=422, detail="Assignment has no campaign to send to.")

    if body.version_id:
        version = await session.get(AssignmentVersionORM, body.version_id)
        if version is None or version.assignment_id != assignment_id:
            raise HTTPException(status_code=404, detail=f"Version '{body.version_id}' not found for this assignment.")
    else:
        result = await session.execute(select(AssignmentVersionORM).where(AssignmentVersionORM.assignment_id == assignment_id).order_by(AssignmentVersionORM.version_number.desc()).limit(1))
        version = result.scalar_one()

    is_first_approval = assignment.status != "approved"
    deadline = datetime.now(timezone.utc) + timedelta(days=body.deadline_days)
    assignment.status = "approved"
    assignment.approved_version_id = version.id
    assignment.deadline = deadline
    assignment.sent_at = datetime.now(timezone.utc)

    notified = 0
    if is_first_approval:
        # Only email candidates the FIRST time an assignment is approved for
        # this campaign - re-approving a different version (rollback) just
        # changes which version future evaluation reads, it doesn't
        # re-notify candidates who already received the assignment.
        result = await session.execute(select(CandidateORM).where(CandidateORM.campaign_id == assignment.campaign_id, CandidateORM.status == "shortlisted"))
        selected_candidates = result.scalars().all()
        email_sender = get_email_sender_dep()
        settings = get_settings()
        for candidate in selected_candidates:
            candidate.submission_token = str(uuid.uuid4())
            candidate.status = "assignment_sent"
            if candidate.email:
                assignment_link, _ = candidate_assignment_link(version, candidate.submission_token)
                email_sender.send(
                    to_email=candidate.email, subject=ASSIGNMENT_EMAIL_SUBJECT,
                    body=ASSIGNMENT_EMAIL_BODY.format(
                        candidate_name=candidate.name or "Candidate",
                        deadline=deadline.strftime("%Y-%m-%d %H:%M UTC"),
                        assignment_link=assignment_link,
                        submission_link=f"{settings.frontend_public_url}/submission/{candidate.submission_token}",
                    ),
                )
                notified += 1

        await _log_event(session, assignment.campaign_id, "assignment_approved", f"Assignment v{version.version_number} approved.")
        await _log_event(session, assignment.campaign_id, "assignment_sent", f"Assignment sent to {notified} candidates.")
    else:
        await _log_event(session, assignment.campaign_id, "assignment_rolled_back", f"Approved version changed to v{version.version_number} (rollback).")

    await session.commit()
    logger.info("Assignment %s approved (v%d), %d candidates notified (0 LLM calls)", assignment_id, version.version_number, notified)
    return AssignmentSendResult(assignment_id=assignment_id, approved_version_number=version.version_number, candidates_notified=notified, deadline=deadline)


@router.post("/campaigns/{campaign_id}/deadline/extend", response_model=AssignmentResponse)
async def extend_deadline(campaign_id: str, body: ExtendDeadlineRequest, session: AsyncSession = Depends(get_session)) -> AssignmentResponse:
    result = await session.execute(select(AssignmentORM).where(AssignmentORM.campaign_id == campaign_id, AssignmentORM.status == "approved").order_by(AssignmentORM.created_at.desc()).limit(1))
    assignment = result.scalar_one_or_none()
    if assignment is None or assignment.deadline is None:
        raise HTTPException(status_code=404, detail="No approved assignment with a deadline found for this campaign.")

    assignment.deadline = _as_utc(assignment.deadline) + timedelta(days=body.additional_days)
    await session.commit()
    await session.refresh(assignment)
    return await _to_response(session, assignment)
