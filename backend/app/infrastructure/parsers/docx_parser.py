"""
DOCX parser implementing the DocumentParser port.

python-docx is synchronous and CPU/IO bound on file reads, so we run it
in a thread pool via asyncio.to_thread to keep the FastAPI event loop
unblocked for concurrent requests.
"""

from __future__ import annotations

import asyncio

from docx import Document

from app.core.exceptions import DocumentParsingError
from app.core.logging import get_logger

logger = get_logger(__name__)


class DocxParser:
    """Extracts paragraph and table text from a .docx file, in document order."""

    async def extract_text(self, file_path: str) -> str:
        try:
            return await asyncio.to_thread(self._extract_sync, file_path)
        except DocumentParsingError:
            raise
        except Exception as e:  # python-docx raises various low-level errors
            logger.exception("Failed to parse DOCX: %s", file_path)
            raise DocumentParsingError(f"Could not read DOCX file: {e}") from e

    @staticmethod
    def _extract_sync(file_path: str) -> str:
        document = Document(file_path)
        parts: list[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        text = "\n".join(parts)
        if not text.strip():
            raise DocumentParsingError(f"No extractable text found in {file_path}")
        return text
